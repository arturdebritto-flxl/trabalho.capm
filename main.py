import argparse
import logging
import sys

import pandas as pd
import numpy as np
import yfinance as yf
import matplotlib.pyplot as plt
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ============================================================
# 1. PARÂMETROS DO PROJETO
# ============================================================

# Identificação do autor.
AUTOR_NOME = "Artur Alessio de Britto"
AUTOR_MATRICULA = "241021696"

# A Embraer utilizava EMBR3 até 31/10/2025.
# A partir de 03/11/2025, passou a utilizar EMBJ3.
TICKER_ACAO_ANTIGO = "EMBR3"
TICKER_ACAO_NOVO = "EMBJ3"
TICKER_MERCADO = "^BVSP"

DATA_INICIO = "2021-01-01"
DATA_INICIO_DOWNLOAD = "2020-12-01"
DATA_TROCA_TICKER = "2025-11-03"
DATA_FIM_ANALISE = "2025-12-31"
# O parâmetro end do yfinance é exclusivo.
# Para incluir dezembro/2025, o download vai até 01/01/2026.
DATA_FIM_DOWNLOAD = "2026-01-01"

# Valor informado na fonte consultada pelo autor.
# Este beta é mantido apenas como referência externa para comparação.
BETA_PUBLICADO_REFERENCIA = 0.82
FONTE_BETA = (
    "STOCK ANALYSIS. Embraer S.A. (BVMF: EMBJ3): Statistics. "
    "Beta: 0.82. Consulta em: 22 maio 2026."
)
URL_FONTE_BETA = "https://stockanalysis.com/quote/bvmf/EMBJ3/statistics/"

# Caminhos absolutos baseados na pasta em que este arquivo main.py está salvo.
BASE_DIR = Path(__file__).resolve().parent
PASTA_DADOS = BASE_DIR / "dados"
PASTA_COTACOES_B3 = PASTA_DADOS / "cotacoes_b3"
PASTA_GRAFICOS = BASE_DIR / "graficos"
PASTA_RELATORIO = BASE_DIR / "relatorio"
PASTA_IMAGENS = BASE_DIR / "imagens"
PASTA_DADOS_STREAMLIT = PASTA_DADOS / "streamlit"

MESES_ESPERADOS = 60  # 5 anos de análise, mensal

logger = logging.getLogger(__name__)


# ============================================================
# 2. COLETA E PREPARAÇÃO DOS DADOS
# ============================================================

def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Executa a análise CAPM da Embraer e gera saída Excel/Word/Gráficos/Streamlit."
    )
    grupo_verbosity = parser.add_mutually_exclusive_group()
    grupo_verbosity.add_argument(
        "--quiet",
        action="store_true",
        help="Minimiza a saída do console."
    )
    grupo_verbosity.add_argument(
        "--debug",
        action="store_true",
        help="Exibe mensagens de depuração detalhadas."
    )
    parser.add_argument(
        "--yes",
        action="store_true",
        help="Executa sem solicitar confirmação e sobrescreve saídas existentes."
    )
    return parser.parse_args()


def configure_logging(quiet: bool, debug: bool) -> None:
    if quiet:
        level = logging.WARNING
    elif debug:
        level = logging.DEBUG
    else:
        level = logging.INFO

    logging.basicConfig(
        level=level,
        format="%(asctime)s - %(levelname)s - %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S"
    )


def pedir_confirmacao() -> bool:
    resposta = input(
        "O script gerará/atualizará arquivos em dados/, graficos/ e relatorio/. Continuar? [s/N]: "
    ).strip().lower()
    return resposta in {"s", "sim"}

def baixar_selic_mensal(data_inicio: str, data_fim: str) -> pd.DataFrame:
    """Baixa a Selic acumulada no mês (série BCB 4390) e converte % para decimal."""
    data_inicio_bcb = datetime.strptime(data_inicio, "%Y-%m-%d").strftime("%d/%m/%Y")
    data_fim_bcb = datetime.strptime(data_fim, "%Y-%m-%d").strftime("%d/%m/%Y")

    url = (
        "https://api.bcb.gov.br/dados/serie/bcdata.sgs.4390/dados"
        f"?formato=json&dataInicial={data_inicio_bcb}&dataFinal={data_fim_bcb}"
    )

    selic = pd.read_json(url)

    if selic.empty:
        raise ValueError("A série da Selic veio vazia. Verifique a conexão e as datas.")

    selic["data"] = pd.to_datetime(selic["data"], dayfirst=True)
    selic["valor"] = pd.to_numeric(selic["valor"], errors="coerce")
    selic = selic.rename(columns={"valor": "selic_mensal"}).set_index("data")
    selic["selic_mensal"] = selic["selic_mensal"] / 100
    selic = selic.resample("ME").last().dropna()

    return selic


def extrair_fechamento(dados: pd.DataFrame, ticker: str) -> pd.Series:
    """Extrai a coluna Close independentemente do formato de colunas retornado pelo yfinance."""
    if dados.empty:
        raise ValueError(f"Nenhum dado foi baixado para {ticker}.")

    if isinstance(dados.columns, pd.MultiIndex):
        fechamento = dados["Close"]
        if isinstance(fechamento, pd.DataFrame):
            fechamento = fechamento.iloc[:, 0]
    else:
        fechamento = dados["Close"]

    return fechamento.astype(float)


def baixar_fechamento_ajustado(
    ticker: str,
    data_inicio: str,
    data_fim: str,
    permitir_vazio: bool = False
) -> pd.Series:
    """Baixa o fechamento ajustado diário de um ticker."""
    dados = yf.download(
        ticker,
        start=data_inicio,
        end=data_fim,
        interval="1d",
        auto_adjust=True,
        progress=False
    )

    if dados.empty:
        if permitir_vazio:
            return pd.Series(dtype="float64", name=ticker)

        raise ValueError(
            f"Nenhum dado foi baixado para {ticker}. "
            "O ticker pode ter sido descontinuado ou não estar disponível na fonte."
        )
    fechamento = extrair_fechamento(dados, ticker)
    fechamento.name = ticker
    return fechamento

def ler_cotahist_b3(caminho_zip: Path) -> pd.DataFrame:
    """
    Lê um arquivo anual COTAHIST da B3 em blocos menores
    e mantém apenas as cotações da Embraer.
    """

    larguras = [
        2, 8, 2, 12, 3, 12, 10, 3, 4,
        13, 13, 13, 13, 13, 13, 13,
        5, 18, 18, 13, 1, 8, 7, 13, 12, 3
    ]

    nomes_colunas = [
        "TIPREG", "DATA", "CODBDI", "CODNEG", "TPMERC",
        "NOMRES", "ESPECI", "PRAZOT", "MODREF",
        "PREABE", "PREMAX", "PREMIN", "PREMED", "PREULT",
        "PREOFC", "PREOFV", "TOTNEG", "QUATOT", "VOLTOT",
        "PREEXE", "INDOPC", "DATVEN", "FATCOT", "PTOEXE",
        "CODISI", "DISMES"
    ]

    colunas_necessarias = [
        "TIPREG", "DATA", "CODNEG", "TPMERC", "PREULT"
    ]

    if not caminho_zip.exists():
        raise FileNotFoundError(
            f"Arquivo da B3 não encontrado: {caminho_zip}"
        )

    partes_embraer = []

    with ZipFile(caminho_zip) as arquivo_zip:
        nomes_txt = [
            nome for nome in arquivo_zip.namelist()
            if nome.upper().endswith(".TXT")
        ]

        if not nomes_txt:
            raise ValueError(
                f"Nenhum arquivo TXT foi encontrado dentro de {caminho_zip.name}."
            )

        nome_txt = nomes_txt[0]

        with arquivo_zip.open(nome_txt) as arquivo_txt:
            leitor_blocos = pd.read_fwf(
                arquivo_txt,
                widths=larguras,
                names=nomes_colunas,
                usecols=colunas_necessarias,
                dtype=str,
                encoding="latin-1",
                chunksize=100000
            )

            for bloco in leitor_blocos:
                bloco["CODNEG"] = bloco["CODNEG"].str.strip()
                bloco["TPMERC"] = bloco["TPMERC"].str.strip()

                bloco_embraer = bloco[
                    (bloco["TIPREG"] == "01")
                    & (bloco["TPMERC"] == "010")
                    & (bloco["CODNEG"].isin([TICKER_ACAO_ANTIGO, TICKER_ACAO_NOVO]))
                ].copy()

                if not bloco_embraer.empty:
                    partes_embraer.append(bloco_embraer)

    if not partes_embraer:
        return pd.DataFrame(
            columns=["TIPREG", "DATA", "CODNEG", "TPMERC", "PREULT"]
        )

    dados = pd.concat(partes_embraer, ignore_index=True)

    dados["DATA"] = pd.to_datetime(
        dados["DATA"],
        format="%Y%m%d",
        errors="coerce"
    )

    dados["PREULT"] = (
        pd.to_numeric(dados["PREULT"], errors="coerce") / 100
    )

    dados = dados.dropna(subset=["DATA", "PREULT"])

    return dados

def carregar_precos_embraer_b3() -> pd.Series:
    """
    Carrega os preços diários da Embraer a partir dos arquivos históricos da B3.
    """

    bases_anuais = []

    for ano in range(2020, 2026):
        caminho_zip = PASTA_COTACOES_B3 / f"COTAHIST_A{ano}.ZIP"

        logger.info(f"Lendo arquivo da B3 de {ano}: {caminho_zip.name}")

        dados_ano = ler_cotahist_b3(caminho_zip)
        logger.debug(
            "Arquivo de %s concluído. Registros da Embraer encontrados: %d",
            ano,
            len(dados_ano),
        )

        if not dados_ano.empty:
            bases_anuais.append(dados_ano)

    if not bases_anuais:
        raise ValueError(
            "Nenhuma cotação da Embraer foi obtida pelos arquivos COTAHIST da B3. "
            "Verifique os arquivos ZIP na pasta dados/cotacoes_b3/."
        )

    cotacoes = pd.concat(bases_anuais, ignore_index=True)

    preco_acao = (
        cotacoes.sort_values("DATA")
        .set_index("DATA")["PREULT"]
        .rename("preco_acao")
    )

    return preco_acao

def baixar_base_precos_continua() -> tuple[pd.DataFrame, str]:
    """
    Monta a série de preços da Embraer e do Ibovespa.
    A Embraer é obtida pela B3 e o Ibovespa pelo yfinance.
    """

    preco_acao = carregar_precos_embraer_b3()
    preco_acao = preco_acao.rename("preco_acao")

    metodo_serie = (
        "Série de preços de fechamento da Embraer obtida nos arquivos "
        "COTAHIST da B3, combinando EMBR3 e EMBJ3. "
        "Os preços da B3 não são ajustados por dividendos, bonificações "
        "ou outros proventos, o que constitui limitação metodológica da análise."
    )

    preco_mercado = baixar_fechamento_ajustado(
        TICKER_MERCADO,
        DATA_INICIO_DOWNLOAD,
        DATA_FIM_DOWNLOAD
    ).rename("preco_mercado")

    precos = pd.concat(
        [preco_acao, preco_mercado],
        axis=1,
        sort=False
    ).dropna()

    if precos.empty:
        raise ValueError("A base conjunta de preços ficou vazia.")

    return precos, metodo_serie


def calcular_retornos_mensais(
    precos: pd.DataFrame
) -> tuple[pd.DataFrame, pd.DataFrame]:
    """Converte preços diários em preços mensais e calcula retornos mensais simples."""
    precos_mensais = precos.resample("ME").last().dropna()
    retornos = precos_mensais.pct_change().dropna()

    retornos = retornos.rename(columns={
        "preco_acao": "retorno_acao",
        "preco_mercado": "retorno_mercado"
    })
    # Remove dezembro/2020 da análise e mantém jan/2021 a dez/2025.
    retornos = retornos.loc[DATA_INICIO:DATA_FIM_ANALISE]
    
    logger.debug("COLUNAS DOS PREÇOS MENSAIS: %s", precos_mensais.columns.tolist())
    logger.debug("COLUNAS DOS RETORNOS: %s", retornos.columns.tolist())

    return precos_mensais, retornos

def calcular_beta(retornos: pd.DataFrame) -> tuple[float, float, float]:
    """
    Calcula o beta da ação em relação ao mercado:
    beta = covariância dos retornos / variância dos retornos do mercado.
    """
    base_beta = retornos[["retorno_acao", "retorno_mercado"]].dropna()

    if len(base_beta) != MESES_ESPERADOS:
        raise ValueError(
            f"Esperavam-se {MESES_ESPERADOS} retornos mensais para calcular o beta, "
            f"mas foram obtidos {len(base_beta)}."
        )

    covariancia = base_beta["retorno_acao"].cov(base_beta["retorno_mercado"])
    variancia_mercado = base_beta["retorno_mercado"].var()

    if np.isclose(variancia_mercado, 0.0):
        raise ValueError(
            "A variância do mercado é zero; não é possível calcular o beta."
        )

    beta = covariancia / variancia_mercado

    return float(beta), float(covariancia), float(variancia_mercado)


def preparar_base_capm(
    retornos: pd.DataFrame,
    selic_mensal: pd.DataFrame
) -> pd.DataFrame:
    """Junta retornos e Selic e calcula retornos excedentes para a regressão."""
    dados_capm = retornos.join(selic_mensal, how="inner").dropna()

    dados_capm["excesso_acao"] = (
        dados_capm["retorno_acao"] - dados_capm["selic_mensal"]
    )
    dados_capm["excesso_mercado"] = (
        dados_capm["retorno_mercado"] - dados_capm["selic_mensal"]
    )

    return dados_capm


def calcular_regressao(dados_capm: pd.DataFrame) -> tuple[float, float, float]:
    """
    Estima a regressão empírica:
    (Ri - Rf) = alpha + beta_estimado * (Rm - Rf) + erro.
    O beta estimado é exibido como diagnóstico complementar.
    O CAPM final utiliza o beta calculado pela covariância/variância dos retornos mensais.
    """
    x = dados_capm["excesso_mercado"]
    y = dados_capm["excesso_acao"]

    beta_estimado, alpha = np.polyfit(x, y, 1)
    correlacao = y.corr(x)
    r_quadrado = correlacao ** 2

    return float(beta_estimado), float(alpha), float(r_quadrado)


# ============================================================
# 3. FORMATAÇÃO E RECOMENDAÇÃO
# ============================================================

def formatar_percentual(valor: float) -> str:
    return f"{valor * 100:.2f}%"


def formatar_decimal(valor: float) -> str:
    return f"{valor:.4f}"


def interpretar_beta(beta: float) -> str:
    if beta > 1:
        return "superior ao mercado"
    if beta < 1:
        return "inferior ao mercado"
    return "semelhante ao mercado"


def intensidade_beta(beta: float) -> str:
    if beta > 1:
        return "mais intensas que as do Ibovespa"
    if beta < 1:
        return "menos intensas que as do Ibovespa"
    return "semelhantes às do Ibovespa"


def classificar_recomendacao(
    retorno_medio_acao: float,
    retorno_esperado_capm: float
) -> tuple[str, float]:
    """
    Compara diretamente retorno observado e retorno exigido pelo CAPM.
    Não adiciona margem arbitrária, pois a Selic já participa do CAPM.
    """
    diferenca = retorno_medio_acao - retorno_esperado_capm

    if diferenca > 0:
        return "Compra", diferenca
    if diferenca < 0:
        return "Venda", diferenca
    return "Manutenção", diferenca


def texto_recomendacao(
    recomendacao: str,
    retorno_medio_acao: float,
    retorno_esperado_capm: float
) -> str:
    if recomendacao == "Compra":
        return (
            f"Como o retorno médio observado da ação, de "
            f"{formatar_percentual(retorno_medio_acao)}, foi superior ao retorno "
            f"exigido pelo CAPM, de {formatar_percentual(retorno_esperado_capm)}, "
            "a recomendação resultante da análise é Compra."
        )

    if recomendacao == "Venda":
        return (
            f"Como o retorno médio observado da ação, de "
            f"{formatar_percentual(retorno_medio_acao)}, foi inferior ao retorno "
            f"exigido pelo CAPM, de {formatar_percentual(retorno_esperado_capm)}, "
            "a recomendação resultante da análise é Venda."
        )

    return (
        f"Como o retorno médio observado da ação, de "
        f"{formatar_percentual(retorno_medio_acao)}, coincidiu com o retorno "
        f"exigido pelo CAPM, de {formatar_percentual(retorno_esperado_capm)}, "
        "a recomendação resultante da análise é Manutenção."
    )


def criar_resumo(
    beta_usado_capm: float,
    beta_publicado_referencia: float,
    covariancia: float,
    variancia_mercado: float,
    beta_estimado: float,
    alpha: float,
    r_quadrado: float,
    retorno_medio_acao: float,
    retorno_medio_mercado: float,
    taxa_livre_risco: float,
    retorno_esperado_capm: float,
    diferenca: float,
    recomendacao: str,
    quantidade_meses: int,
    metodo_serie: str
) -> pd.DataFrame:
    """Cria a tabela-resumo já formatada para relatório e planilha."""
    indicadores = [
        "Empresa analisada",
        "Ticker atual",
        "Índice de mercado",
        "Período analisado",
        "Frequência",
        "Quantidade de retornos mensais",
        "Beta calculado por covariância/variância usado no CAPM",
        "Beta publicado de 5 anos apenas para referência",
        "Covariância entre retorno da ação e retorno do mercado",
        "Variância dos retornos do mercado",
        "Beta da regressão com retornos excedentes",
        "Alfa da regressão",
        "R² da regressão",
        "Retorno médio da ação ao mês",
        "Retorno médio do mercado ao mês",
        "Taxa livre de risco média ao mês",
        "Retorno esperado pelo CAPM ao mês",
        "Diferença entre retorno observado e CAPM",
        "Recomendação de investimento",
        "Construção da série da ação",
        "Fonte do beta publicado externo",
    ]
    valores = [
        "Embraer S.A.",
        TICKER_ACAO_NOVO,
        TICKER_MERCADO,
        f"{DATA_INICIO} a {DATA_FIM_ANALISE}",
        "Mensal",
        str(quantidade_meses),
        formatar_decimal(beta_usado_capm),
        formatar_decimal(beta_publicado_referencia),
        f"{covariancia:.8f}",
        f"{variancia_mercado:.8f}",
        formatar_decimal(beta_estimado),
        formatar_decimal(alpha),
        formatar_decimal(r_quadrado),
        formatar_percentual(retorno_medio_acao),
        formatar_percentual(retorno_medio_mercado),
        formatar_percentual(taxa_livre_risco),
        formatar_percentual(retorno_esperado_capm),
        formatar_percentual(diferenca),
        recomendacao,
        metodo_serie,
        FONTE_BETA,
    ]

    if len(indicadores) != len(valores):
        raise ValueError(
            "A tabela de resumo possui quantidade diferente de indicadores e valores."
        )

    return pd.DataFrame({"Indicador": indicadores, "Valor": valores})


# ============================================================
# 4. GRÁFICOS E EXCEL
# ============================================================

def gerar_graficos(
    dados_capm: pd.DataFrame,
    beta_estimado: float,
    alpha: float
) -> tuple[Path, Path]:
    """Gera os gráficos utilizados no Excel e no Word."""
    PASTA_GRAFICOS.mkdir(exist_ok=True)

    caminho_retornos = PASTA_GRAFICOS / "retornos_mensais_embraer_ibovespa.png"
    caminho_regressao = PASTA_GRAFICOS / "regressao_capm.png"

    plt.figure(figsize=(10, 6))
    plt.plot(dados_capm.index, dados_capm["retorno_acao"], label="Embraer")
    plt.plot(dados_capm.index, dados_capm["retorno_mercado"], label="Ibovespa")
    plt.axhline(0, linestyle="--")
    plt.title("Retornos mensais: Embraer x Ibovespa")
    plt.xlabel("Data")
    plt.ylabel("Retorno mensal")
    plt.legend()
    plt.tight_layout()
    plt.savefig(caminho_retornos, dpi=200)
    plt.close()

    x = dados_capm["excesso_mercado"]
    y = dados_capm["excesso_acao"]
    linha_regressao = alpha + beta_estimado * x

    ordem = np.argsort(x.to_numpy())

    plt.figure(figsize=(8, 6))
    plt.scatter(x, y, label="Observações mensais")
    plt.plot(x.iloc[ordem], linha_regressao.iloc[ordem], label="Linha de regressão")
    plt.title("Regressão CAPM: Embraer x Ibovespa")
    plt.xlabel("Retorno excedente do Ibovespa")
    plt.ylabel("Retorno excedente da Embraer")
    plt.legend()
    plt.tight_layout()
    plt.savefig(caminho_regressao, dpi=200)
    plt.close()

    return caminho_retornos, caminho_regressao


def exportar_excel_organizado(
    resumo: pd.DataFrame,
    dados_capm: pd.DataFrame,
    precos_mensais: pd.DataFrame,
    caminho_grafico_retornos: Path,
    caminho_grafico_regressao: Path
) -> Path:
    """Exporta um único arquivo Excel organizado, com resumo, base e preços."""
    PASTA_DADOS.mkdir(exist_ok=True)
    caminho_excel = PASTA_DADOS / "analise_capm_embraer.xlsx"

    base_mensal = dados_capm.reset_index().rename(columns={
        "Date": "Data",
        "data": "Data",
        "retorno_acao": "Retorno Embraer",
        "retorno_mercado": "Retorno Ibovespa",
        "selic_mensal": "Selic mensal",
        "excesso_acao": "Excesso Embraer",
        "excesso_mercado": "Excesso Ibovespa"
    })
    base_mensal = base_mensal.rename(columns={base_mensal.columns[0]: "Data"})

    precos_exportar = precos_mensais.reset_index()
    precos_exportar = precos_exportar.rename(columns={
        precos_exportar.columns[0]: "Data",
        "preco_acao": "Preço Embraer",
        "preco_mercado": "Ibovespa"
    })

    with pd.ExcelWriter(caminho_excel, engine="xlsxwriter", datetime_format="dd/mm/yyyy") as writer:
        resumo.to_excel(writer, sheet_name="Resumo", index=False, startrow=2)
        base_mensal.to_excel(writer, sheet_name="Base mensal", index=False)
        precos_exportar.to_excel(writer, sheet_name="Preços mensais", index=False)

        workbook = writer.book
        ws_resumo = writer.sheets["Resumo"]
        ws_base = writer.sheets["Base mensal"]
        ws_precos = writer.sheets["Preços mensais"]

        formato_titulo = workbook.add_format({
            "bold": True,
            "font_size": 16,
            "align": "center",
            "valign": "vcenter",
            "font_color": "white",
            "bg_color": "#1F4E78"
        })
        formato_cabecalho = workbook.add_format({
            "bold": True,
            "font_color": "white",
            "bg_color": "#1F4E78",
            "border": 1
        })
        formato_texto = workbook.add_format({
            "border": 1,
            "text_wrap": True,
            "valign": "top"
        })
        formato_percentual = workbook.add_format({
            "num_format": "0.00%",
            "border": 1
        })
        formato_numero = workbook.add_format({
            "num_format": "0.0000",
            "border": 1
        })
        formato_data = workbook.add_format({
            "num_format": "dd/mm/yyyy",
            "border": 1
        })
        formato_recomendacao = workbook.add_format({
            "bold": True,
            "bg_color": "#FFF2CC",
            "border": 1
        })

        ws_resumo.merge_range("A1:B1", "ANÁLISE CAPM - EMBRAER S.A.", formato_titulo)
        ws_resumo.set_row(0, 28)
        ws_resumo.set_column("A:A", 48)
        ws_resumo.set_column("B:B", 68)
        ws_resumo.freeze_panes(3, 0)

        # Reaplica bordas e realça a recomendação.
        for linha in range(len(resumo)):
            linha_excel = linha + 3
            indicador = resumo.iloc[linha, 0]
            valor = resumo.iloc[linha, 1]
            ws_resumo.write(linha_excel, 0, indicador, formato_texto)
            formato_valor = formato_recomendacao if indicador == "Recomendação de investimento" else formato_texto
            ws_resumo.write(linha_excel, 1, valor, formato_valor)

        if caminho_grafico_retornos.exists():
            ws_resumo.insert_image("D3", str(caminho_grafico_retornos), {"x_scale": 0.52, "y_scale": 0.52})
        if caminho_grafico_regressao.exists():
            ws_resumo.insert_image("D25", str(caminho_grafico_regressao), {"x_scale": 0.52, "y_scale": 0.52})

        ws_base.set_row(0, 22, formato_cabecalho)
        ws_base.set_column("A:A", 14, formato_data)
        ws_base.set_column("B:G", 20, formato_percentual)
        ws_base.freeze_panes(1, 0)
        ws_base.autofilter(0, 0, len(base_mensal), len(base_mensal.columns) - 1)

        ws_precos.set_row(0, 22, formato_cabecalho)
        ws_precos.set_column("A:A", 14, formato_data)
        ws_precos.set_column("B:C", 18, formato_numero)
        ws_precos.freeze_panes(1, 0)
        ws_precos.autofilter(0, 0, len(precos_exportar), len(precos_exportar.columns) - 1)

    return caminho_excel


def criar_campos_calendario(datas: pd.Series) -> pd.DataFrame:
    """Cria campos mensais padronizados para o painel Streamlit."""
    nomes_meses = {
        1: "janeiro",
        2: "fevereiro",
        3: "março",
        4: "abril",
        5: "maio",
        6: "junho",
        7: "julho",
        8: "agosto",
        9: "setembro",
        10: "outubro",
        11: "novembro",
        12: "dezembro",
    }

    calendario = pd.DataFrame({"Data": pd.to_datetime(datas)})
    calendario["Ano"] = calendario["Data"].dt.year
    calendario["MesNumero"] = calendario["Data"].dt.month
    calendario["MesNome"] = calendario["MesNumero"].map(nomes_meses)
    calendario["AnoMes"] = calendario["Data"].dt.strftime("%Y-%m")
    calendario["OrdemAnoMes"] = calendario["Ano"] * 100 + calendario["MesNumero"]
    calendario["Trimestre"] = "T" + calendario["Data"].dt.quarter.astype(str)
    return calendario


def annualizar_taxa(taxa_mensal: float) -> float:
    """Converte uma taxa média mensal em taxa anual equivalente."""
    return (1 + taxa_mensal) ** 12 - 1


def exportar_dados_streamlit(
    dados_capm: pd.DataFrame,
    precos_mensais: pd.DataFrame,
    resumo: pd.DataFrame,
    beta_usado_capm: float,
    beta_publicado_referencia: float,
    covariancia: float,
    variancia_mercado: float,
    beta_estimado: float,
    alpha: float,
    r_quadrado: float,
    retorno_medio_acao: float,
    retorno_medio_mercado: float,
    taxa_livre_risco: float,
    retorno_esperado_capm: float,
    diferenca: float,
    recomendacao: str,
    metodo_serie: str
) -> Path:
    """Exporta bases CSV para o painel Streamlit complementar."""
    del resumo

    PASTA_DADOS_STREAMLIT.mkdir(parents=True, exist_ok=True)

    base = dados_capm.join(
        precos_mensais[["preco_acao", "preco_mercado"]],
        how="left"
    ).copy()
    base = base.reset_index().rename(columns={base.index.name or "index": "Data"})
    base = base.rename(columns={base.columns[0]: "Data"})
    base["Data"] = pd.to_datetime(base["Data"])

    if len(base) != MESES_ESPERADOS:
        raise ValueError(
            f"A base Streamlit deve conter {MESES_ESPERADOS} meses, "
            f"mas contém {len(base)}."
        )

    campos_calendario = criar_campos_calendario(base["Data"])
    fato_capm = pd.concat(
        [base[["Data"]], campos_calendario.drop(columns=["Data", "Trimestre"])],
        axis=1
    )
    fato_capm["Empresa"] = "Embraer S.A."
    fato_capm["TickerAtual"] = TICKER_ACAO_NOVO
    fato_capm["IndiceMercado"] = TICKER_MERCADO
    fato_capm["PrecoAcao"] = base["preco_acao"]
    fato_capm["PrecoMercado"] = base["preco_mercado"]
    fato_capm["RetornoAcao"] = base["retorno_acao"]
    fato_capm["RetornoMercado"] = base["retorno_mercado"]
    fato_capm["SelicMensal"] = base["selic_mensal"]
    fato_capm["ExcessoAcao"] = base["excesso_acao"]
    fato_capm["ExcessoMercado"] = base["excesso_mercado"]
    fato_capm["RetornoAcumuladoAcao"] = (1 + fato_capm["RetornoAcao"]).cumprod() - 1
    fato_capm["RetornoAcumuladoMercado"] = (1 + fato_capm["RetornoMercado"]).cumprod() - 1

    colunas_fato = [
        "Data",
        "Ano",
        "MesNumero",
        "MesNome",
        "AnoMes",
        "OrdemAnoMes",
        "Empresa",
        "TickerAtual",
        "IndiceMercado",
        "PrecoAcao",
        "PrecoMercado",
        "RetornoAcao",
        "RetornoMercado",
        "SelicMensal",
        "ExcessoAcao",
        "ExcessoMercado",
        "RetornoAcumuladoAcao",
        "RetornoAcumuladoMercado",
    ]
    fato_capm[colunas_fato].to_csv(
        PASTA_DADOS_STREAMLIT / "fato_capm_mensal.csv",
        index=False,
        encoding="utf-8",
        date_format="%Y-%m-%d",
    )

    indicadores = [
        ("Risco", "Beta calculado usado no CAPM", beta_usado_capm, formatar_decimal(beta_usado_capm), "decimal", "Calculado no Python por covariância/variância."),
        ("Referência", "Beta publicado apenas para referência", beta_publicado_referencia, formatar_decimal(beta_publicado_referencia), "decimal", "Não utilizado no CAPM final."),
        ("Risco", "Covariância ação-mercado", covariancia, f"{covariancia:.8f}", "decimal", "Covariância entre retorno da ação e retorno do mercado."),
        ("Risco", "Variância do mercado", variancia_mercado, f"{variancia_mercado:.8f}", "decimal", "Variância dos retornos do Ibovespa."),
        ("Regressão", "Beta da regressão com retornos excedentes", beta_estimado, formatar_decimal(beta_estimado), "decimal", "Diagnóstico complementar."),
        ("Regressão", "Alfa", alpha, formatar_decimal(alpha), "decimal", "Alfa da regressão mensal."),
        ("Regressão", "R²", r_quadrado, formatar_decimal(r_quadrado), "decimal", "Coeficiente de determinação da regressão."),
        ("Retorno", "Retorno médio mensal da ação", retorno_medio_acao, formatar_percentual(retorno_medio_acao), "% ao mês", "Média dos 60 retornos mensais da Embraer."),
        ("Retorno", "Retorno médio mensal do mercado", retorno_medio_mercado, formatar_percentual(retorno_medio_mercado), "% ao mês", "Média dos 60 retornos mensais do Ibovespa."),
        ("Retorno", "Selic média mensal", taxa_livre_risco, formatar_percentual(taxa_livre_risco), "% ao mês", "Média da Selic mensal, série BCB 4390."),
        ("CAPM", "Retorno esperado CAPM mensal", retorno_esperado_capm, formatar_percentual(retorno_esperado_capm), "% ao mês", "Calculado no Python com o beta oficial da análise."),
        ("CAPM", "Diferença observado menos CAPM mensal", diferenca, formatar_percentual(diferenca), "% ao mês", "Retorno médio observado menos retorno esperado pelo CAPM."),
        ("Retorno", "Retorno anualizado da ação", annualizar_taxa(retorno_medio_acao), formatar_percentual(annualizar_taxa(retorno_medio_acao)), "% ao ano", "Taxa anual equivalente a partir da média mensal."),
        ("Retorno", "Retorno anualizado do mercado", annualizar_taxa(retorno_medio_mercado), formatar_percentual(annualizar_taxa(retorno_medio_mercado)), "% ao ano", "Taxa anual equivalente a partir da média mensal."),
        ("Retorno", "Selic anualizada", annualizar_taxa(taxa_livre_risco), formatar_percentual(annualizar_taxa(taxa_livre_risco)), "% ao ano", "Taxa anual equivalente a partir da média mensal."),
        ("CAPM", "Retorno CAPM anualizado", annualizar_taxa(retorno_esperado_capm), formatar_percentual(annualizar_taxa(retorno_esperado_capm)), "% ao ano", "Taxa anual equivalente a partir do CAPM mensal."),
        ("Resultado", "Recomendação final", np.nan, recomendacao, "texto", "Classificação calculada pelo Python."),
        ("Base", "Quantidade de retornos mensais", len(dados_capm), str(len(dados_capm)), "meses", "Jan/2021 a dez/2025."),
    ]
    pd.DataFrame(
        indicadores,
        columns=["Grupo", "Indicador", "ValorNumerico", "ValorFormatado", "Unidade", "Observacao"],
    ).to_csv(
        PASTA_DADOS_STREAMLIT / "resumo_indicadores.csv",
        index=False,
        encoding="utf-8",
    )

    calendario_datas = pd.date_range(
        start=DATA_INICIO_DOWNLOAD,
        end=DATA_FIM_ANALISE,
        freq="ME",
    )
    criar_campos_calendario(pd.Series(calendario_datas)).to_csv(
        PASTA_DADOS_STREAMLIT / "dimensao_calendario.csv",
        index=False,
        encoding="utf-8",
        date_format="%Y-%m-%d",
    )

    metodologia_fontes = pd.DataFrame([
        {
            "Item": "Embraer",
            "Descricao": "Preços de fechamento da Embraer combinando EMBR3 e EMBJ3 no mercado à vista.",
            "Fonte": "Arquivos históricos COTAHIST da B3 em dados/cotacoes_b3/.",
            "Limitacao": "Preços da B3 não ajustados por dividendos, bonificações ou outros proventos.",
        },
        {
            "Item": "Ibovespa",
            "Descricao": "Proxy de mercado da análise CAPM.",
            "Fonte": "Yahoo Finance, ticker ^BVSP, via yfinance.",
            "Limitacao": "Sujeito à disponibilidade e eventuais revisões do provedor.",
        },
        {
            "Item": "Selic",
            "Descricao": "Taxa livre de risco mensal.",
            "Fonte": "Banco Central do Brasil, série SGS 4390.",
            "Limitacao": "Usada como proxy mensal da taxa livre de risco.",
        },
        {
            "Item": "Beta",
            "Descricao": "Beta principal calculado por covariância(retorno da ação, retorno do mercado) / variância(retorno do mercado).",
            "Fonte": "Cálculo próprio em Python com os 60 retornos mensais.",
            "Limitacao": "Beta externo publicado é mantido apenas como referência comparativa.",
        },
        {
            "Item": "Aviso",
            "Descricao": "Aplicação acadêmica do CAPM à Embraer.",
            "Fonte": "Projeto Python e relatório Word do repositório.",
            "Limitacao": "Não constitui recomendação financeira definitiva.",
        },
        {
            "Item": "Construção da série",
            "Descricao": metodo_serie,
            "Fonte": "Processamento Python do projeto.",
            "Limitacao": "Dezembro/2020 é usado apenas como preço-base para o retorno de janeiro/2021.",
        },
    ])
    metodologia_fontes.to_csv(
        PASTA_DADOS_STREAMLIT / "metodologia_fontes.csv",
        index=False,
        encoding="utf-8",
    )

    return PASTA_DADOS_STREAMLIT


# ============================================================
# 5. RELATÓRIO WORD
# ============================================================

def adicionar_paragrafo_centralizado(
    documento: Document,
    texto: str = "",
    negrito: bool = False,
    tamanho: int = 12,
    espaco_depois: int = 0
):
    """Adiciona parágrafo centralizado com formatação consistente."""
    paragrafo = documento.add_paragraph()
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragrafo.paragraph_format.space_after = Pt(espaco_depois)

    run = paragrafo.add_run(texto)
    run.bold = negrito
    run.font.name = "Times New Roman"
    run.font.size = Pt(tamanho)
    return paragrafo


def localizar_logo_unb() -> Path | None:
    """Localiza a logo original da UnB salva na pasta imagens."""
    for nome_arquivo in (
        "logo_unb.png",
        "logo_unb.jpg",
        "logo_unb.jpeg",
    ):
        caminho = PASTA_IMAGENS / nome_arquivo
        if caminho.exists():
            return caminho

    return None


def adicionar_capa_unb(documento: Document):
    """Insere a capa no padrão do template da nota técnica."""
    section = documento.sections[0]
    section.top_margin = Cm(3)
    section.left_margin = Cm(3)
    section.bottom_margin = Cm(2)
    section.right_margin = Cm(2)

    logo_path = localizar_logo_unb()

    if logo_path is not None:
        p_logo = documento.add_paragraph()
        p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_logo.paragraph_format.space_after = Pt(8)

        run_logo = p_logo.add_run()
        run_logo.add_picture(str(logo_path), width=Inches(1.2))
    else:
        adicionar_paragrafo_centralizado(
            documento,
            "[INSERIR LOGO DA UnB EM imagens/logo_unb.png]",
            tamanho=10,
            espaco_depois=8
        )
        print(
            "Aviso: logo da UnB não encontrada. "
            "Verifique o arquivo imagens/logo_unb.png."
        )

    adicionar_paragrafo_centralizado(documento, "Universidade de Brasília - UnB", negrito=True)
    adicionar_paragrafo_centralizado(
        documento,
        "Faculdade de Economia, Administração, Contabilidade e Gestão de Políticas Públicas - FACE"
    )
    adicionar_paragrafo_centralizado(documento, "Departamento de Ciências Contábeis e Atuariais - CCA")
    adicionar_paragrafo_centralizado(documento, "Bacharelado em Ciências Contábeis", espaco_depois=22)

    adicionar_paragrafo_centralizado(
        documento,
        "Nota Técnica nº ___/Disciplina MQC                                      Data: _______",
        espaco_depois=26
    )
    adicionar_paragrafo_centralizado(
        documento,
        "NOTA TÉCNICA - RECOMENDAÇÃO DE COMPRA DE AÇÃO VIA CAPM",
        negrito=True,
        tamanho=14,
        espaco_depois=22
    )
    adicionar_paragrafo_centralizado(documento, "Empresa analisada: Embraer S.A.")
    adicionar_paragrafo_centralizado(documento, "Setor: Aeronáutico / bens industriais")
    adicionar_paragrafo_centralizado(
        documento,
        f"Período de análise: {DATA_INICIO} a {DATA_FIM_ANALISE} (freq.: mensal)",
        espaco_depois=28
    )
    adicionar_paragrafo_centralizado(documento, AUTOR_NOME, negrito=True)
    adicionar_paragrafo_centralizado(documento, f"Matrícula: {AUTOR_MATRICULA}", espaco_depois=34)
    adicionar_paragrafo_centralizado(documento, "Brasília - DF")
    adicionar_paragrafo_centralizado(documento, "2026")

    documento.add_page_break()


def adicionar_pagina_autoridades(documento: Document):
    """Insere a página institucional conforme o template fornecido."""
    autoridades = [
        ("Professora Doutora Márcia Abrahão Moura", "Reitora da Universidade de Brasília"),
        ("Professor Doutor Enrique Huelva Unternbäumen", "Vice-Reitor da Universidade de Brasília"),
        ("Professor Doutor Sérgio Antônio Andrade de Freitas", "Decano de Ensino de Graduação"),
        (
            "Professor Doutor Eduardo Tadeu Vieira",
            "Diretor da Faculdade de Economia, Administração, Contabilidade e Gestão de Políticas Públicas"
        ),
        (
            "Professor Doutor Paulo César de Melo Mendes",
            "Chefe do Departamento de Ciências Contábeis e Atuariais"
        ),
        (
            "Professor Doutor Alex Laquis Resende",
            "Coordenador de Graduação do curso de Ciências Contábeis - Diurno"
        ),
        (
            "Professor Mestre Elivânio Geraldo de Andrade",
            "Coordenador de Graduação do curso de Ciências Contábeis - Noturno"
        )
    ]

    for nome, cargo in autoridades:
        adicionar_paragrafo_centralizado(documento, nome)
        adicionar_paragrafo_centralizado(documento, cargo, espaco_depois=16)

    documento.add_page_break()


def gerar_relatorio_word(
    resumo: pd.DataFrame,
    beta_usado_capm: float,
    beta_publicado_referencia: float,
    covariancia: float,
    variancia_mercado: float,
    beta_estimado: float,
    alpha: float,
    r_quadrado: float,
    retorno_medio_acao: float,
    retorno_medio_mercado: float,
    taxa_livre_risco: float,
    retorno_esperado_capm: float,
    diferenca: float,
    recomendacao: str,
    metodo_serie: str,
    caminho_grafico_retornos: Path,
    caminho_grafico_regressao: Path
) -> Path:
    """Gera a nota técnica Word completa com resultados, recomendação e gráficos."""
    PASTA_RELATORIO.mkdir(exist_ok=True)
    caminho_saida = PASTA_RELATORIO / "nota_tecnica_capm_embraer.docx"

    documento = Document()
    estilo_normal = documento.styles["Normal"]
    estilo_normal.font.name = "Times New Roman"
    estilo_normal.font.size = Pt(12)

    adicionar_capa_unb(documento)
    adicionar_pagina_autoridades(documento)

    documento.add_heading("1 SUMÁRIO EXECUTIVO", level=1)
    documento.add_paragraph(
        f"Esta nota técnica analisa a ação EMBJ3, da Embraer S.A., por meio do "
        f"Capital Asset Pricing Model (CAPM), utilizando dados mensais do período "
        f"de {DATA_INICIO} a {DATA_FIM_ANALISE}."
    )
    documento.add_paragraph(
    f"O beta calculado a partir dos retornos mensais da ação e do Ibovespa, "
    f"utilizado no CAPM, foi de {formatar_decimal(beta_usado_capm)}. "
    f"Para comparação, o beta publicado de cinco anos consultado externamente "
    f"foi de {formatar_decimal(beta_publicado_referencia)}. "
    f"O retorno médio observado da ação foi de "
    f"{formatar_percentual(retorno_medio_acao)} ao mês, enquanto o "
    f"retorno exigido pelo CAPM foi de "
    f"{formatar_percentual(retorno_esperado_capm)} ao mês. "
    f"A recomendação final é: {recomendacao}."
)

    documento.add_heading("2 CONTEXTO E OBJETIVO", level=1)
    documento.add_paragraph(
        "A Embraer S.A. é uma companhia brasileira do setor aeronáutico, com "
        "atuação em aviação comercial, executiva, defesa e serviços. A ação "
        "atualmente negociada na B3 é a EMBJ3."
    )
    documento.add_paragraph(
        "O objetivo desta nota técnica é estimar o retorno exigido da ação por "
        "meio do CAPM e compará-lo com o retorno médio observado, produzindo uma "
        "recomendação de Compra, Venda ou Manutenção."
    )

    documento.add_heading("3 METODOLOGIA", level=1)
    documento.add_heading("3.1 Capital Asset Pricing Model (CAPM)", level=2)
    documento.add_paragraph(
        "O CAPM relaciona o retorno exigido do ativo ao risco sistemático. "
        "A expressão aplicada foi:"
    )
    adicionar_paragrafo_centralizado(documento, "E(Ri) = Rf + βi [E(Rm) - Rf]", negrito=True, espaco_depois=8)
    documento.add_paragraph(
        "Em que E(Ri) é o retorno exigido da ação; Rf é a taxa livre de risco; "
        "βi é o beta empregado no modelo; e E(Rm) é o retorno médio do mercado."
    )
    documento.add_paragraph(
        f"Para o cálculo final do CAPM foi utilizado o beta calculado pela covariância/variância "
        f"dos retornos mensais, conforme apresentado nos resultados ({formatar_decimal(beta_usado_capm)}). "
        "A regressão sobre a base mensal é apresentada separadamente como verificação empírica e diagnóstico."
    )

    documento.add_heading("3.2 Estratégia empírica", level=2)
    documento.add_paragraph(
        f"Foram coletados preços de fechamento da Embraer nos arquivos COTAHIST "
        f"da B3, combinando EMBR3 e EMBJ3, e dados do Ibovespa por meio da "
        f"biblioteca yfinance. {metodo_serie}"
    )
    documento.add_paragraph(
        "Os preços diários foram convertidos em preços mensais pelo último pregão "
        "de cada mês. Os retornos mensais foram calculados pela variação percentual. "
        "A Selic mensal, obtida pela série 4390 do Banco Central do Brasil, foi "
        "utilizada como taxa livre de risco."
    )

    documento.add_heading("3.2.1 Regressão Linear", level=2)
    adicionar_paragrafo_centralizado(
        documento,
        "(Ri - Rf) = α + β (Rm - Rf) + ε",
        negrito=True,
        espaco_depois=8
    )
    documento.add_paragraph(
        f"A regressão dos retornos excedentes resultou em beta estimado de "
        f"{formatar_decimal(beta_estimado)}, alfa de {formatar_decimal(alpha)} "
        f"e R² de {formatar_decimal(r_quadrado)}. Esse beta empírico é informado "
        "como análise complementar e verificação; o CAPM final utiliza o beta calculado pelo programa."
    )

    documento.add_heading("3.2.2 Custo de capital e retorno esperado", level=2)
    documento.add_paragraph(
        "O retorno esperado calculado pelo CAPM representa a remuneração mínima "
        "exigida pelo investidor diante do risco sistemático adotado na análise."
    )

    documento.add_heading("3.3 Coleta e tratamento dos dados", level=2)
    documento.add_paragraph(
        "A leitura dos arquivos COTAHIST da B3, a coleta do Ibovespa com yfinance "
        "e o tratamento dos dados foram automatizados em Python com pandas; os "
        "gráficos foram produzidos com matplotlib; e as saídas foram geradas em "
        "Excel e Word."
    )

    documento.add_heading("4 RESULTADOS", level=1)
    documento.add_heading("4.1 Risco sistemático", level=2)
    documento.add_paragraph(
        f"O beta utilizado no CAPM foi {formatar_decimal(beta_usado_capm)}, "
        f"indicando risco sistemático {interpretar_beta(beta_usado_capm)}. "
        f"Com esse beta, as oscilações esperadas da ação tendem a ser "
        f"{intensidade_beta(beta_usado_capm)}."
    )

    documento.add_heading("4.2 Custo de capital - retorno esperado", level=2)
    tabela = documento.add_table(rows=1, cols=2)
    tabela.style = "Table Grid"
    tabela.rows[0].cells[0].text = "Indicador"
    tabela.rows[0].cells[1].text = "Valor"
    for _, linha in resumo.iterrows():
        celulas = tabela.add_row().cells
        celulas[0].text = str(linha["Indicador"])
        celulas[1].text = str(linha["Valor"])

    documento.add_paragraph(
        f"O retorno esperado pelo CAPM foi de "
        f"{formatar_percentual(retorno_esperado_capm)} ao mês, enquanto o retorno "
        f"médio observado foi de {formatar_percentual(retorno_medio_acao)} ao mês. "
        f"A diferença foi de {formatar_percentual(diferenca)} ao mês."
    )

    documento.add_heading("4.3 Gráficos da análise", level=2)
    if caminho_grafico_retornos.exists():
        documento.add_picture(str(caminho_grafico_retornos), width=Inches(5.8))
        adicionar_paragrafo_centralizado(
            documento,
            "Figura 1 - Retornos mensais da Embraer e do Ibovespa.",
            tamanho=10,
            espaco_depois=8
        )
    if caminho_grafico_regressao.exists():
        documento.add_picture(str(caminho_grafico_regressao), width=Inches(5.8))
        adicionar_paragrafo_centralizado(
            documento,
            "Figura 2 - Regressão entre retornos excedentes da Embraer e do Ibovespa.",
            tamanho=10,
            espaco_depois=8
        )

    documento.add_heading("4.4 Riscos, limitações e notas contábeis", level=2)
    documento.add_paragraph(
        "A análise pelo CAPM considera o risco sistemático e depende das premissas "
        "adotadas, especialmente a taxa livre de risco, a proxy de mercado e o beta. "
        "O beta publicado disponível externamente pode ter metodologia distinta da regressão "
        "mensal apresentada como análise complementar; aqui o beta calculado pelo programa foi usado no CAPM."
    )
    documento.add_paragraph(
        "Para a Embraer, também devem ser considerados fatores específicos, como "
        "variação cambial, demanda internacional, custos de produção, juros, cadeia "
        "de suprimentos e ciclos econômicos globais."
    )
    documento.add_paragraph(
        "Os preços de fechamento da Embraer extraídos dos arquivos COTAHIST da B3 "
        "não são ajustados por dividendos, bonificações ou outros proventos. Essa "
        "característica constitui limitação metodológica da análise."
    )

    documento.add_heading("5 CONCLUSÃO E RECOMENDAÇÕES", level=1)
    documento.add_paragraph(
        f"O CAPM foi aplicado com o beta calculado pela covariância/variância dos retornos "
        f"(valor utilizado: {formatar_decimal(beta_usado_capm)}). O retorno exigido calculado foi de "
        f"{formatar_percentual(retorno_esperado_capm)} ao mês e o retorno médio "
        f"observado da ação foi de {formatar_percentual(retorno_medio_acao)} ao mês."
    )
    documento.add_paragraph(
        texto_recomendacao(recomendacao, retorno_medio_acao, retorno_esperado_capm)
    )

    documento.add_heading("REFERÊNCIAS", level=1)
    documento.add_paragraph(
        "BANCO CENTRAL DO BRASIL. Sistema Gerenciador de Séries Temporais - "
        "Taxa Selic acumulada no mês, série 4390."
    )
    documento.add_paragraph(
        "B3. Arquivos históricos COTAHIST de cotações da Embraer."
    )
    documento.add_paragraph(
        "YAHOO FINANCE. Dados históricos de preços do Ibovespa (^BVSP)."
    )
    documento.add_paragraph(FONTE_BETA)
    documento.add_paragraph(URL_FONTE_BETA)

    documento.save(caminho_saida)
    return caminho_saida


# ============================================================
# 6. FUNÇÃO PRINCIPAL
# ============================================================

def main(yes: bool = False) -> None:
    """Executa a análise, gera uma planilha Excel, gráficos e o relatório Word."""
    if not yes:
        logger.info(
            "Aviso: esta execução irá gerar ou sobrescrever saídas em dados/, graficos/ e relatorio/."
        )
        if not pedir_confirmacao():
            logger.warning("Execução cancelada pelo usuário.")
            return
    else:
        logger.info("Modo automático ativado (--yes): executando sem confirmação.")

    PASTA_DADOS.mkdir(exist_ok=True)
    PASTA_GRAFICOS.mkdir(exist_ok=True)
    PASTA_RELATORIO.mkdir(exist_ok=True)
    PASTA_IMAGENS.mkdir(exist_ok=True)
    PASTA_DADOS_STREAMLIT.mkdir(parents=True, exist_ok=True)

    selic_mensal = baixar_selic_mensal(DATA_INICIO, DATA_FIM_ANALISE)
    precos, metodo_serie = baixar_base_precos_continua()
    precos_mensais, retornos = calcular_retornos_mensais(precos)
    
    beta_calculado, covariancia, variancia_mercado = calcular_beta(retornos)
    
    logger.debug("VERIFICAÇÃO DOS RETORNOS:\n%s\n%s", retornos.head(), retornos.tail())
    logger.info("Quantidade de retornos mensais: %d", len(retornos))

    logger.debug(
        "VERIFICAÇÃO DO BETA: covariância=%f, variância=%f, beta=%f",
        covariancia,
        variancia_mercado,
        beta_calculado,
    )
    
    dados_capm = preparar_base_capm(retornos, selic_mensal)

    if len(dados_capm) != MESES_ESPERADOS:
        raise ValueError(
            f"Base de análise incompleta: obtidos {len(dados_capm)} meses, "
            f"mas são necessários exatamente {MESES_ESPERADOS} meses (2021-01 a 2025-12). "
            "Verifique os arquivos COTAHIST em dados/cotacoes_b3/ e a cobertura de dados do Ibovespa."
        )

    beta_estimado, alpha, r_quadrado = calcular_regressao(dados_capm)

    beta_usado_capm = beta_calculado
    retorno_medio_acao = dados_capm["retorno_acao"].mean()
    retorno_medio_mercado = dados_capm["retorno_mercado"].mean()
    taxa_livre_risco = dados_capm["selic_mensal"].mean()

    retorno_esperado_capm = taxa_livre_risco + beta_usado_capm * (
        retorno_medio_mercado - taxa_livre_risco
    )

    recomendacao, diferenca = classificar_recomendacao(
        retorno_medio_acao,
        retorno_esperado_capm
    )

    resumo = criar_resumo(
        beta_usado_capm=beta_usado_capm,
        beta_publicado_referencia=BETA_PUBLICADO_REFERENCIA,
        covariancia=covariancia,
        variancia_mercado=variancia_mercado,
        beta_estimado=beta_estimado,
        alpha=alpha,
        r_quadrado=r_quadrado,
        retorno_medio_acao=retorno_medio_acao,
        retorno_medio_mercado=retorno_medio_mercado,
        taxa_livre_risco=taxa_livre_risco,
        retorno_esperado_capm=retorno_esperado_capm,
        diferenca=diferenca,
        recomendacao=recomendacao,
        quantidade_meses=len(dados_capm),
        metodo_serie=metodo_serie
    )

    caminho_grafico_retornos, caminho_grafico_regressao = gerar_graficos(
        dados_capm,
        beta_estimado,
        alpha
    )

    caminho_excel = exportar_excel_organizado(
        resumo,
        dados_capm,
        precos_mensais,
        caminho_grafico_retornos,
        caminho_grafico_regressao
    )

    caminho_word = gerar_relatorio_word(
        resumo=resumo,
        beta_usado_capm=beta_usado_capm,
        beta_publicado_referencia=BETA_PUBLICADO_REFERENCIA,
        covariancia=covariancia,
        variancia_mercado=variancia_mercado,
        beta_estimado=beta_estimado,
        alpha=alpha,
        r_quadrado=r_quadrado,
        retorno_medio_acao=retorno_medio_acao,
        retorno_medio_mercado=retorno_medio_mercado,
        taxa_livre_risco=taxa_livre_risco,
        retorno_esperado_capm=retorno_esperado_capm,
        diferenca=diferenca,
        recomendacao=recomendacao,
        metodo_serie=metodo_serie,
        caminho_grafico_retornos=caminho_grafico_retornos,
        caminho_grafico_regressao=caminho_grafico_regressao
    )

    caminho_streamlit = exportar_dados_streamlit(
        dados_capm=dados_capm,
        precos_mensais=precos_mensais,
        resumo=resumo,
        beta_usado_capm=beta_usado_capm,
        beta_publicado_referencia=BETA_PUBLICADO_REFERENCIA,
        covariancia=covariancia,
        variancia_mercado=variancia_mercado,
        beta_estimado=beta_estimado,
        alpha=alpha,
        r_quadrado=r_quadrado,
        retorno_medio_acao=retorno_medio_acao,
        retorno_medio_mercado=retorno_medio_mercado,
        taxa_livre_risco=taxa_livre_risco,
        retorno_esperado_capm=retorno_esperado_capm,
        diferenca=diferenca,
        recomendacao=recomendacao,
        metodo_serie=metodo_serie,
    )

    print("\nRESULTADOS PRINCIPAIS")
    print(f"Beta calculado usado no CAPM: {beta_usado_capm:.4f}")
    print(f"Beta publicado externo apenas para referência: {BETA_PUBLICADO_REFERENCIA:.4f}")
    print(f"Covariância ação x mercado: {covariancia:.8f}")
    print(f"Variância dos retornos do mercado: {variancia_mercado:.8f}")
    print(f"Beta da regressão com retornos excedentes: {beta_estimado:.4f}")
    print(f"Alfa da regressão: {alpha:.4f}")
    print(f"R²: {r_quadrado:.4f}")
    print(f"Retorno médio da ação: {formatar_percentual(retorno_medio_acao)} ao mês")
    print(f"Retorno esperado pelo CAPM: {formatar_percentual(retorno_esperado_capm)} ao mês")
    print(f"Recomendação: {recomendacao}")
    print(f"Excel: {caminho_excel}")
    print(f"Word: {caminho_word}")
    print(f"Dados Streamlit: {caminho_streamlit}")


if __name__ == "__main__":
    args = parse_args()
    configure_logging(args.quiet, args.debug)
    try:
        main(yes=args.yes)
    except KeyboardInterrupt:
        logger.warning("Execução interrompida pelo usuário.")
        sys.exit(1)
